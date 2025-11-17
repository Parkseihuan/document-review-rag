"""DOCX file parser"""
from docx import Document
from typing import Optional


class DOCXParser:
    """Parse DOCX files to markdown"""

    @staticmethod
    def parse(file_path: str) -> Optional[str]:
        """
        Parse DOCX file and extract text

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text in markdown format or None if parsing fails
        """
        try:
            doc = Document(file_path)
            markdown_content = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                # Check if paragraph is a heading
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        markdown_content.append(f"{'#' * level_num} {text}\n")
                    except ValueError:
                        markdown_content.append(f"{text}\n")
                else:
                    markdown_content.append(f"{text}\n")

            # Extract tables
            for table in doc.tables:
                markdown_content.append("\n")
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_content.append("| " + " | ".join(cells) + " |")
                    if i == 0:  # Add header separator
                        markdown_content.append("| " + " | ".join(["---"] * len(cells)) + " |")
                markdown_content.append("\n")

            return '\n'.join(markdown_content)

        except Exception as e:
            print(f"Error parsing DOCX {file_path}: {e}")
            return None

    @staticmethod
    def to_markdown(file_path: str, output_path: str) -> bool:
        """
        Parse DOCX and save as markdown file

        Args:
            file_path: Path to DOCX file
            output_path: Path to save markdown file

        Returns:
            True if successful, False otherwise
        """
        content = DOCXParser.parse(file_path)
        if content:
            try:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                return True
            except Exception as e:
                print(f"Error writing markdown file: {e}")
                return False
        return False
